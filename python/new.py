import re
from docx import Document
 
def docx_replace_regex(doc_obj, reg_obj , replace):

    for p in doc_obj.paragraphs:
   
        if reg_obj.search(p.text):
            inline = p.runs
        
            
            for i in range(len(inline)):
                if reg_obj.search(inline[i].text):
                    text = reg_obj.sub(replace, inline[i].text)
                    inline[i].text = text
 
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, reg_obj , replace)
 
   
 
regobj = re.compile("<project_name>")

 
new_word = input('the project name')
filename = "SERVER SETUP DOCUMENTATION.docx"
doc = Document(filename)
ext = '.docx'
filenew = new_word + ext
#a = doc.paragraphs[100].text
#print (a)
#url = re.findall('http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\), ]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', a) 
#print (url)
docx_replace_regex(doc, regobj , new_word)
doc.save(filenew)